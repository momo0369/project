# -*- coding: utf-8 -*-
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime, timedelta
import os
import json
import re
import uuid
import requests
import smtplib
from email.mime.text import MIMEText
from email.header import Header
from werkzeug.utils import secure_filename

app = Flask(__name__, static_folder='.', static_url_path='')
CORS(app)

app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///project_management.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

db = SQLAlchemy(app)

if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])


class User(db.Model):
    __tablename__ = 'user'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    username = db.Column(db.String(50), unique=True, nullable=False)
    password = db.Column(db.String(100), nullable=False)
    name = db.Column(db.String(50), nullable=False)
    email = db.Column(db.String(100))
    phone = db.Column(db.String(20))
    role = db.Column(db.String(20), default='member')
    can_attend_meeting = db.Column(db.Boolean, default=True)
    can_assign_task = db.Column(db.Boolean, default=True)
    created_at = db.Column(db.DateTime, default=datetime.now)
    updated_at = db.Column(db.DateTime, default=datetime.now, onupdate=datetime.now)


class Project(db.Model):
    __tablename__ = 'project'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    name = db.Column(db.String(100), nullable=False)
    code = db.Column(db.String(50), unique=True, nullable=False)
    objective = db.Column(db.Text)
    background = db.Column(db.Text)
    start_date = db.Column(db.Date)
    end_date = db.Column(db.Date)
    manager_id = db.Column(db.Integer, db.ForeignKey('user.id'))
    owner_id = db.Column(db.Integer, db.ForeignKey('user.id'))
    status = db.Column(db.String(20), default='not_started')
    description = db.Column(db.Text)
    project_type = db.Column(db.String(50))
    created_at = db.Column(db.DateTime, default=datetime.now)
    updated_at = db.Column(db.DateTime, default=datetime.now, onupdate=datetime.now)
    manager = db.relationship('User', foreign_keys=[manager_id], backref='managed_projects')
    owner = db.relationship('User', foreign_keys=[owner_id], backref='owned_projects')


class ProjectMember(db.Model):
    __tablename__ = 'project_member'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    project_id = db.Column(db.Integer, db.ForeignKey('project.id'), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    role = db.Column(db.String(20), default='member')
    joined_at = db.Column(db.DateTime, default=datetime.now)
    project = db.relationship('Project', backref='members')
    user = db.relationship('User', backref='project_memberships')


class Milestone(db.Model):
    __tablename__ = 'milestone'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    project_id = db.Column(db.Integer, db.ForeignKey('project.id'), nullable=False)
    name = db.Column(db.String(100), nullable=False)
    objective = db.Column(db.Text)
    planned_date = db.Column(db.Date)
    actual_date = db.Column(db.Date)
    manager_id = db.Column(db.Integer, db.ForeignKey('user.id'))
    status = db.Column(db.String(20), default='not_started')
    description = db.Column(db.Text)
    need_meeting = db.Column(db.Boolean, default=False)
    expected_meeting_date = db.Column(db.Date)
    meeting_participants = db.Column(db.Text)
    meeting_objective = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=datetime.now)
    updated_at = db.Column(db.DateTime, default=datetime.now, onupdate=datetime.now)
    project = db.relationship('Project', backref='milestones')
    manager = db.relationship('User', backref='managed_milestones')


class Task(db.Model):
    __tablename__ = 'task'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    project_id = db.Column(db.Integer, db.ForeignKey('project.id'), nullable=False)
    milestone_id = db.Column(db.Integer, db.ForeignKey('milestone.id'))
    name = db.Column(db.String(100), nullable=False)
    objective = db.Column(db.Text)
    description = db.Column(db.Text)
    manager_id = db.Column(db.Integer, db.ForeignKey('user.id'))
    participants = db.Column(db.Text)
    start_date = db.Column(db.Date)
    due_date = db.Column(db.Date)
    priority = db.Column(db.String(20), default='medium')
    status = db.Column(db.String(20), default='not_started')
    need_meeting = db.Column(db.Boolean, default=False)
    meeting_objective = db.Column(db.Text)
    suggested_participants = db.Column(db.Text)
    source = db.Column(db.String(20), default='manual')
    meeting_id = db.Column(db.Integer, db.ForeignKey('meeting.id'))
    meeting_reminder_created = db.Column(db.Boolean, default=False)
    meeting_invitation_created = db.Column(db.Boolean, default=False)
    suggested_meeting_date = db.Column(db.Date)
    suggested_meeting_topic = db.Column(db.String(200))
    created_at = db.Column(db.DateTime, default=datetime.now)
    updated_at = db.Column(db.DateTime, default=datetime.now, onupdate=datetime.now)
    project = db.relationship('Project', backref='tasks')
    milestone = db.relationship('Milestone', backref='tasks')
    manager = db.relationship('User', backref='assigned_tasks')
    source_meeting = db.relationship('Meeting', foreign_keys=[meeting_id], backref='generated_tasks')


class Meeting(db.Model):
    __tablename__ = 'meeting'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    project_id = db.Column(db.Integer, db.ForeignKey('project.id'), nullable=False)
    title = db.Column(db.String(200), nullable=False)
    objective = db.Column(db.Text)
    meeting_date = db.Column(db.DateTime)
    location = db.Column(db.String(200))
    meeting_type = db.Column(db.String(50))
    organizer_id = db.Column(db.Integer, db.ForeignKey('user.id'))
    description = db.Column(db.Text)
    source = db.Column(db.String(20), default='manual')
    source_task_id = db.Column(db.Integer, db.ForeignKey('task.id'))
    source_milestone_id = db.Column(db.Integer, db.ForeignKey('milestone.id'))
    status = db.Column(db.String(20), default='scheduled')
    minutes_uploaded = db.Column(db.Boolean, default=False)
    tasks_generated = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=datetime.now)
    updated_at = db.Column(db.DateTime, default=datetime.now, onupdate=datetime.now)
    project = db.relationship('Project', backref='meetings')
    organizer = db.relationship('User', backref='organized_meetings')
    source_task = db.relationship('Task', foreign_keys=[source_task_id])
    source_milestone = db.relationship('Milestone', foreign_keys=[source_milestone_id])


class MeetingAttendee(db.Model):
    __tablename__ = 'meeting_attendee'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    meeting_id = db.Column(db.Integer, db.ForeignKey('meeting.id'), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    is_required = db.Column(db.Boolean, default=True)
    meeting = db.relationship('Meeting', backref='attendees')
    user = db.relationship('User', backref='meeting_attendances')


class MeetingMinutes(db.Model):
    __tablename__ = 'meeting_minutes'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    meeting_id = db.Column(db.Integer, db.ForeignKey('meeting.id'), nullable=False)
    file_path = db.Column(db.String(500))
    file_name = db.Column(db.String(200))
    content = db.Column(db.Text)
    conclusion = db.Column(db.Text)
    parse_status = db.Column(db.String(20), default='uploaded')
    parse_error = db.Column(db.Text)
    generated_task_count = db.Column(db.Integer, default=0)
    parsed_at = db.Column(db.DateTime)
    created_at = db.Column(db.DateTime, default=datetime.now)
    meeting = db.relationship('Meeting', backref='minutes')


class ActionItem(db.Model):
    __tablename__ = 'action_item'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    minutes_id = db.Column(db.Integer, db.ForeignKey('meeting_minutes.id'), nullable=False)
    title = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text)
    assignee_id = db.Column(db.Integer, db.ForeignKey('user.id'))
    participants = db.Column(db.Text)
    due_date = db.Column(db.Date)
    priority = db.Column(db.String(20), default='medium')
    status = db.Column(db.String(20), default='pending')
    original_text = db.Column(db.Text)
    confidence = db.Column(db.Float, default=0.8)
    task_id = db.Column(db.Integer, db.ForeignKey('task.id'))
    created_at = db.Column(db.DateTime, default=datetime.now)
    minutes = db.relationship('MeetingMinutes', backref='action_items')
    assignee = db.relationship('User', backref='action_items')
    task = db.relationship('Task', backref='source_action_item')


class Notification(db.Model):
    __tablename__ = 'notification'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    title = db.Column(db.String(200), nullable=False)
    type = db.Column(db.String(50), nullable=False)
    project_id = db.Column(db.Integer, db.ForeignKey('project.id'))
    related_type = db.Column(db.String(50))
    related_id = db.Column(db.Integer)
    content = db.Column(db.Text)
    is_read = db.Column(db.Boolean, default=False)
    is_handled = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=datetime.now)
    user = db.relationship('User', backref='notifications')
    project = db.relationship('Project', backref='notifications')


class RuleConfig(db.Model):
    __tablename__ = 'rule_config'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    rule_key = db.Column(db.String(100), unique=True, nullable=False)
    rule_name = db.Column(db.String(100), nullable=False)
    rule_value = db.Column(db.Text)
    description = db.Column(db.Text)
    is_active = db.Column(db.Boolean, default=True)
    created_at = db.Column(db.DateTime, default=datetime.now)
    updated_at = db.Column(db.DateTime, default=datetime.now, onupdate=datetime.now)


class OperationLog(db.Model):
    __tablename__ = 'operation_log'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'))
    action = db.Column(db.String(100), nullable=False)
    target_type = db.Column(db.String(50))
    target_id = db.Column(db.Integer)
    detail = db.Column(db.Text)
    ip_address = db.Column(db.String(50))
    created_at = db.Column(db.DateTime, default=datetime.now)
    user = db.relationship('User', backref='operation_logs')


def serialize_model(obj, exclude=None):
    if obj is None:
        return None
    result = {}
    for column in obj.__table__.columns:
        if exclude and column.name in exclude:
            continue
        value = getattr(obj, column.name)
        if isinstance(value, datetime):
            value = value.strftime('%Y-%m-%d %H:%M:%S')
        elif isinstance(value, type(None)):
            value = None
        else:
            from datetime import date
            if isinstance(value, date):
                value = value.strftime('%Y-%m-%d')
        result[column.name] = value
    return result


def get_task_display_status(task, today=None):
    if today is None:
        today = datetime.now().date()
    if task.status == 'completed':
        return 'completed'
    if task.start_date and today < task.start_date:
        return 'not_started'
    if task.due_date and today > task.due_date:
        return 'overdue'
    return 'in_progress'


def parse_id_list(value):
    if not value:
        return []
    values = value if isinstance(value, list) else str(value).split(',')
    result = []
    for item in values:
        try:
            parsed = int(str(item).strip())
        except (TypeError, ValueError):
            continue
        if parsed not in result:
            result.append(parsed)
    return result


def parse_optional_date(value):
    if not value:
        return None
    return datetime.strptime(value, '%Y-%m-%d').date()


def parse_meeting_datetime(value):
    if not value:
        return None
    formats = ['%Y-%m-%d %H:%M', '%Y-%m-%dT%H:%M', '%Y-%m-%dT%H:%M:%S', '%Y-%m-%d']
    for fmt in formats:
        try:
            return datetime.strptime(value, fmt)
        except ValueError:
            continue
    raise ValueError(f'日期格式错误: {value}')


def get_task_meeting_datetime(task, now=None):
    if now is None:
        now = datetime.now()
    base_date = task.suggested_meeting_date or now.date()
    if isinstance(base_date, datetime):
        base_date = base_date.date()
    return datetime.combine(base_date, datetime.min.time()).replace(
        hour=10, minute=0, second=0, microsecond=0
    )


def build_task_meeting_defaults(task, now=None):
    meeting_date = get_task_meeting_datetime(task, now)
    attendee_ids = []
    if task.manager_id:
        attendee_ids.append(task.manager_id)
    for participant_id in parse_id_list(task.participants):
        if participant_id not in attendee_ids:
            attendee_ids.append(participant_id)
    return {
        'project_id': task.project_id,
        'title': f'{task.name}推进会议',
        'meeting_date': meeting_date.strftime('%Y-%m-%dT%H:%M'),
        'meeting_date_display': meeting_date.strftime('%Y-%m-%d %H:%M'),
        'objective': task.objective,
        'description': task.description,
        'attendee_ids': attendee_ids
    }


def build_milestone_meeting_defaults(milestone, now=None):
    if now is None:
        now = datetime.now()
    base_date = milestone.expected_meeting_date or now.date()
    if isinstance(base_date, datetime):
        base_date = base_date.date()
    meeting_date = datetime.combine(base_date, datetime.min.time()).replace(
        hour=10, minute=0, second=0, microsecond=0
    )
    attendee_ids = []
    for user_id in [milestone.manager_id, milestone.project.manager_id if milestone.project else None, milestone.project.owner_id if milestone.project else None]:
        if user_id and user_id not in attendee_ids:
            attendee_ids.append(user_id)
    return {
        'project_id': milestone.project_id,
        'title': f'{milestone.name}评审会议',
        'meeting_date': meeting_date.strftime('%Y-%m-%dT%H:%M'),
        'meeting_date_display': meeting_date.strftime('%Y-%m-%d %H:%M'),
        'objective': milestone.meeting_objective or milestone.objective,
        'description': milestone.description,
        'attendee_ids': attendee_ids
    }


def build_meeting_reminders(limit=None, user_id=None):
    reminders_with_sort = []

    tasks_need_meeting = Task.query.filter_by(
        need_meeting=True, meeting_invitation_created=False
    ).all()
    for task in tasks_need_meeting:
        if user_id and task.manager_id != user_id:
            continue
        meeting_defaults = build_task_meeting_defaults(task)
        reminder = {
            'type': 'task',
            'source_id': task.id,
            'source_name': task.name,
            'project_id': task.project_id,
            'project_name': task.project.name if task.project else None,
            'reason': '任务标记需要会议但未创建会议邀请',
            'suggested_date': meeting_defaults['meeting_date_display'],
            'suggested_topic': meeting_defaults['title'],
            'objective': meeting_defaults['objective'],
            'description': meeting_defaults['description'],
            'attendee_ids': meeting_defaults['attendee_ids'],
            'meeting_defaults': meeting_defaults
        }
        reminders_with_sort.append((get_task_meeting_datetime(task), reminder))

    milestones_need_meeting = Milestone.query.filter_by(need_meeting=True).all()
    for milestone in milestones_need_meeting:
        has_meeting = Meeting.query.filter_by(source_milestone_id=milestone.id).first()
        if has_meeting:
            continue
        if user_id:
            allowed_user_ids = {
                milestone.manager_id,
                milestone.project.manager_id if milestone.project else None,
                milestone.project.owner_id if milestone.project else None
            }
            if user_id not in allowed_user_ids:
                continue
        suggested_at = None
        if milestone.expected_meeting_date:
            suggested_at = datetime.combine(milestone.expected_meeting_date, datetime.min.time())
        meeting_defaults = build_milestone_meeting_defaults(milestone)
        reminder = {
            'type': 'milestone',
            'source_id': milestone.id,
            'source_name': milestone.name,
            'project_id': milestone.project_id,
            'project_name': milestone.project.name if milestone.project else None,
            'reason': '里程碑需要阶段评审会议',
            'suggested_date': meeting_defaults['meeting_date_display'],
            'suggested_topic': meeting_defaults['title'],
            'suggested_participants': milestone.meeting_participants,
            'objective': meeting_defaults['objective'],
            'description': meeting_defaults['description'],
            'attendee_ids': meeting_defaults['attendee_ids'],
            'meeting_defaults': meeting_defaults
        }
        reminders_with_sort.append((suggested_at, reminder))

    reminders_with_sort.sort(
        key=lambda item: (
            item[0] is None,
            item[0] or datetime.max,
            item[1]['type'],
            item[1]['source_id']
        )
    )
    reminders = [item[1] for item in reminders_with_sort]
    if limit is not None:
        reminders = reminders[:limit]
    return reminders


def get_meeting_type_label(meeting_type):
    label_map = {
        'offline': '线下会议',
        'online': '线上会议'
    }
    return label_map.get(meeting_type, meeting_type or '-')


def create_meeting_invitation_notifications(meeting, attendee_ids):
    project_name = meeting.project.name if meeting.project else '-'
    meeting_time = meeting.meeting_date.strftime('%Y-%m-%d %H:%M') if meeting.meeting_date else '-'
    meeting_type = get_meeting_type_label(meeting.meeting_type)
    location = meeting.location or '-'
    objective = meeting.objective or '-'
    description = meeting.description or '-'
    attendees = []
    group_content = (
        f'您有一个新的会议待参加，请查收：\n'
        f'所属项目：{project_name}\n'
        f'会议主题：{meeting.title}\n'
        f'会议时间：{meeting_time}\n'
        f'会议方式：{meeting_type}\n'
        f'会议地方或链接：{location}\n'
        f'会议目标：{objective}\n'
        f'会议描述：{description}'
    )
    for attendee_id in parse_id_list(attendee_ids):
        attendee = User.query.get(attendee_id)
        if not attendee:
            continue
        attendees.append(attendee)
        content = (
            f'{attendee.name}您好，您有一个新的会议待参加，请查收：\n'
            f'所属项目：{project_name}\n'
            f'会议主题：{meeting.title}\n'
            f'会议时间：{meeting_time}\n'
            f'会议方式：{meeting_type}\n'
            f'会议地方或链接：{location}\n'
            f'会议目标：{objective}\n'
            f'会议描述：{description}'
        )
        create_notification(
            user_id=attendee.id,
            title=f'会议邀请：{meeting.title}',
            type='meeting_invitation',
            project_id=meeting.project_id,
            related_type='meeting',
            related_id=meeting.id,
            content=content,
            send_feishu=False
        )
    if attendees:
        send_feishu_notification(attendees, group_content)


def get_rule_value(rule_key, default=None):
    env_map = {
        'enable_email_notification': 'ENABLE_EMAIL_NOTIFICATION',
        'enable_feishu_notification': 'ENABLE_FEISHU_NOTIFICATION',
        'feishu_webhook_url': 'FEISHU_WEBHOOK_URL',
        'smtp_server': 'SMTP_SERVER',
        'smtp_port': 'SMTP_PORT',
        'smtp_sender': 'SMTP_SENDER',
        'smtp_password': 'SMTP_PASSWORD',
        'minutes_llm_chat_url': 'MEETING_MINUTES_LLM_CHAT_URL',
        'minutes_llm_model': 'MEETING_MINUTES_LLM_MODEL',
        'minutes_llm_temperature': 'MEETING_MINUTES_LLM_TEMPERATURE',
        'minutes_llm_timeout_seconds': 'MEETING_MINUTES_LLM_TIMEOUT_SECONDS'
    }
    env_key = env_map.get(rule_key)
    if env_key and os.getenv(env_key):
        return os.getenv(env_key)
    rule = RuleConfig.query.filter_by(rule_key=rule_key).first()
    if rule and rule.rule_value not in [None, '']:
        return rule.rule_value
    return default


def is_rule_enabled(rule_key, default=False):
    value = str(get_rule_value(rule_key, 'true' if default else 'false')).strip().lower()
    return value in ['1', 'true', 'yes', 'on']


def send_email_notification(recipient_email, subject, content):
    if not recipient_email or not is_rule_enabled('enable_email_notification', default=False):
        return False
    smtp_server = get_rule_value('smtp_server', '')
    smtp_port = parse_optional_int(get_rule_value('smtp_port', '465'))
    smtp_sender = get_rule_value('smtp_sender', '')
    smtp_password = get_rule_value('smtp_password', '')
    if not smtp_server or not smtp_port or not smtp_sender or not smtp_password:
        print('=== 邮件通知跳过：SMTP 配置不完整 ===')
        return False
    msg = MIMEText(content or '', 'plain', 'utf-8')
    msg['From'] = smtp_sender
    msg['To'] = recipient_email
    msg['Subject'] = Header(subject, 'utf-8')
    try:
        with smtplib.SMTP_SSL(smtp_server, smtp_port, timeout=30) as server:
            server.login(smtp_sender, smtp_password)
            server.sendmail(smtp_sender, [recipient_email], msg.as_string())
        return True
    except Exception as exc:
        print(f'=== 邮件发送失败: {recipient_email} / {exc} ===')
        return False


def send_feishu_notification(users, content):
    if not is_rule_enabled('enable_feishu_notification', default=False):
        return False
    webhook_url = get_rule_value('feishu_webhook_url', '')
    if not webhook_url:
        print('=== 飞书通知跳过：Webhook 未配置 ===')
        return False
    unique_users = []
    unique_ids = set()
    for user in users or []:
        if not user or user.id in unique_ids:
            continue
        unique_users.append(user)
        unique_ids.add(user.id)
    mention_text = ' '.join([f'@{user.name}' for user in unique_users]).strip()
    text = f'{mention_text}\n{content}' if mention_text else (content or '')
    payload = {
        'msg_type': 'text',
        'content': {
            'text': text
        }
    }
    try:
        response = requests.post(webhook_url, json=payload, timeout=20)
        response.raise_for_status()
        return True
    except Exception as exc:
        print(f'=== 飞书通知发送失败: {exc} ===')
        return False


def create_notification(user_id, title, type, project_id=None, related_type=None, related_id=None, content=None, send_feishu=True):
    notification = Notification(
        user_id=user_id,
        title=title,
        type=type,
        project_id=project_id,
        related_type=related_type,
        related_id=related_id,
        content=content
    )
    db.session.add(notification)
    user = User.query.get(user_id) if user_id else None
    if user and user.email:
        send_email_notification(user.email, title, content or title)
    if send_feishu and user:
        send_feishu_notification([user], content or title)
    return notification


def read_text_file(file_path):
    encodings = ['utf-8', 'utf-8-sig', 'gbk', 'gb2312', 'utf-16', 'latin1']
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                return file.read()
        except UnicodeDecodeError:
            continue
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
        return file.read()


def extract_text_from_pdf(file_path):
    import fitz

    text_parts = []
    with fitz.open(file_path) as pdf_doc:
        for page in pdf_doc:
            text_parts.append(page.get_text("text"))
    return '\n'.join(text_parts)


def extract_text_from_docx(file_path):
    from docx import Document

    doc = Document(file_path)
    text_parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_parts.append(' | '.join(row_text))
    return '\n'.join(text_parts)


def extract_text_from_doc(file_path):
    import pythoncom
    import win32com.client

    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        word = win32com.client.DispatchEx('Word.Application')
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(os.path.abspath(file_path), ReadOnly=True)
        return doc.Content.Text
    finally:
        if doc is not None:
            doc.Close(False)
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()


def extract_minutes_text(file_path, original_name):
    extension = os.path.splitext(original_name or file_path)[1].lower()
    if extension == '.txt':
        return read_text_file(file_path)
    if extension == '.pdf':
        return extract_text_from_pdf(file_path)
    if extension == '.docx':
        return extract_text_from_docx(file_path)
    if extension == '.doc':
        return extract_text_from_doc(file_path)
    raise ValueError('仅支持 TXT、PDF、DOC、DOCX 格式的会议纪要文件')


def clean_text_content(text):
    if not text:
        return ''
    text = text.replace('\x00', ' ')
    text = re.sub(r'\r\n?', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def build_minutes_project_context(meeting):
    project = meeting.project
    milestones = Milestone.query.filter_by(project_id=project.id).order_by(Milestone.planned_date).all()
    project_members = ProjectMember.query.filter_by(project_id=project.id).all()
    member_map = {}
    for project_member in project_members:
        member_map[project_member.user_id] = project_member.user
    if project.manager_id and project.manager:
        member_map[project.manager_id] = project.manager
    if project.owner_id and project.owner:
        member_map[project.owner_id] = project.owner
    members = list(member_map.values())
    return {
        'project': project,
        'milestones': milestones,
        'members': members
    }


def build_minutes_task_prompt(meeting, minutes_content):
    context = build_minutes_project_context(meeting)
    project = context['project']
    members = context['members']
    milestones = context['milestones']
    member_names = [member.name for member in members]
    milestone_names = [milestone.name for milestone in milestones]
    meeting_time = meeting.meeting_date.strftime('%Y-%m-%d %H:%M') if meeting.meeting_date else ''
    prompt_text = f"""
你是一个专业的项目管理助手。

【项目基本信息】
项目名称：{project.name}
项目目标：{project.objective or ''}
项目描述：{project.description or ''}
项目背景：{project.background or ''}
项目成员：{json.dumps(member_names, ensure_ascii=False)}
项目里程碑：{json.dumps(milestone_names, ensure_ascii=False)}

【会议信息】
会议主题：{meeting.title}
会议时间：{meeting_time}
会议目标：{meeting.objective or ''}
会议描述：{meeting.description or ''}

【会议纪要内容】
{minutes_content}

【任务】
请根据会议纪要，提取其中提到的待办事项，并生成任务列表。
每个任务必须包含以下字段：
- project（所属项目，固定为当前项目名称）
- milestone（所属里程碑，优先从上面的项目里程碑中选择；如果无法明确匹配则返回空字符串）
- task_name（任务名称）
- assignee（负责人，必须是项目成员之一）
- priority（优先级：高 / 中 / 低）
- start_date（开始时间，YYYY-MM-DD）
- end_date（截止时间，YYYY-MM-DD）
- task_goal（任务目标）
- task_desc（任务描述）
- participants（参与人列表，元素必须是项目成员姓名）
- need_meeting（是否需要创建会议，true/false）
- suggested_meeting_date（建议会议日期，YYYY-MM-DD，如不需要则为空字符串）

【输出要求】
- 仅返回标准 JSON 数组
- 不要包含解释、说明或 markdown
"""
    return prompt_text, context


def call_minutes_llm(prompt_text):
    chat_url = get_rule_value('minutes_llm_chat_url', 'http://192.168.200.54:11434/api/chat')
    model = get_rule_value('minutes_llm_model', 'glm-5:cloud')
    temperature = float(get_rule_value('minutes_llm_temperature', '0.2'))
    timeout_seconds = int(float(get_rule_value('minutes_llm_timeout_seconds', '180')))
    payload = {
        'model': model,
        'messages': [
            {'role': 'system', 'content': '你是一个专业的项目管理助手'},
            {'role': 'user', 'content': prompt_text}
        ],
        'stream': False,
        'options': {
            'temperature': temperature
        }
    }
    response = requests.post(chat_url, json=payload, timeout=timeout_seconds)
    response.raise_for_status()
    result = response.json()
    content = (
        result.get('message', {}).get('content')
        or result.get('choices', [{}])[0].get('message', {}).get('content')
        or ''
    )
    if not content:
        raise ValueError('大模型未返回有效内容')
    return content


def extract_json_array(content):
    if not content:
        raise ValueError('大模型返回内容为空')
    cleaned = content.strip()
    if cleaned.startswith('```'):
        cleaned = re.sub(r'^```(?:json)?', '', cleaned).strip()
        cleaned = re.sub(r'```$', '', cleaned).strip()
    start = cleaned.find('[')
    end = cleaned.rfind(']')
    if start == -1 or end == -1 or end < start:
        raise ValueError('未找到有效的 JSON 数组')
    return json.loads(cleaned[start:end + 1])


def normalize_match_key(value):
    if value is None:
        return ''
    return re.sub(r'[\s_\-/（）()【】\[\]：:、，,。.]', '', str(value)).lower()


def match_user_by_name(name, members):
    target = normalize_match_key(name)
    if not target:
        return None
    for member in members:
        if normalize_match_key(member.name) == target:
            return member
    for member in members:
        member_key = normalize_match_key(member.name)
        if target in member_key or member_key in target:
            return member
    return None


def match_milestone_by_name(name, milestones):
    target = normalize_match_key(name)
    if not target:
        return None
    for milestone in milestones:
        if normalize_match_key(milestone.name) == target:
            return milestone
    for milestone in milestones:
        milestone_key = normalize_match_key(milestone.name)
        if target in milestone_key or milestone_key in target:
            return milestone
    return None


def parse_flexible_date(value, fallback=None):
    if not value:
        return fallback
    value = str(value).strip()
    formats = ['%Y-%m-%d', '%Y/%m/%d', '%Y.%m.%d', '%Y年%m月%d日']
    for fmt in formats:
        try:
            return datetime.strptime(value, fmt).date()
        except ValueError:
            continue
    match = re.search(r'(\d{4})\D?(\d{1,2})\D?(\d{1,2})', value)
    if match:
        return datetime(int(match.group(1)), int(match.group(2)), int(match.group(3))).date()
    return fallback


def normalize_priority(value):
    value = str(value or '').strip().lower()
    mapping = {
        '高': 'high',
        '高优先级': 'high',
        'high': 'high',
        'h': 'high',
        '中': 'medium',
        '中优先级': 'medium',
        'medium': 'medium',
        'm': 'medium',
        '低': 'low',
        '低优先级': 'low',
        'low': 'low',
        'l': 'low'
    }
    return mapping.get(value, 'medium')


def normalize_bool(value):
    if isinstance(value, bool):
        return value
    value = str(value or '').strip().lower()
    return value in ['true', '1', 'yes', 'y', '是', '需要']


def normalize_minutes_task_items(task_items, context, meeting_date):
    project = context['project']
    milestones = context['milestones']
    members = context['members']
    normalized_items = []
    errors = []
    for index, item in enumerate(task_items, start=1):
        task_name = (item.get('task_name') or item.get('title') or '').strip()
        assignee_name = (item.get('assignee') or '').strip()
        assignee = match_user_by_name(assignee_name, members)
        if not task_name:
            errors.append(f'第{index}项缺少任务名称')
            continue
        if not assignee:
            errors.append(f'第{index}项负责人无法匹配：{assignee_name or "未填写"}')
            continue
        milestone = match_milestone_by_name(item.get('milestone'), milestones)
        start_date = parse_flexible_date(item.get('start_date'), fallback=meeting_date)
        end_date = parse_flexible_date(item.get('end_date'), fallback=start_date)
        participant_names = item.get('participants') or []
        if isinstance(participant_names, str):
            participant_names = re.split(r'[、,，;；\s]+', participant_names)
        participant_ids = []
        for participant_name in participant_names:
            member = match_user_by_name(participant_name, members)
            if member and member.id not in participant_ids:
                participant_ids.append(member.id)
        if assignee.id not in participant_ids:
            participant_ids.insert(0, assignee.id)
        need_meeting = normalize_bool(item.get('need_meeting'))
        suggested_meeting_date = parse_flexible_date(item.get('suggested_meeting_date'))
        normalized_items.append({
            'project_id': project.id,
            'milestone_id': milestone.id if milestone else None,
            'task_name': task_name,
            'assignee_id': assignee.id,
            'assignee_name': assignee.name,
            'priority': normalize_priority(item.get('priority')),
            'start_date': start_date,
            'end_date': end_date,
            'task_goal': (item.get('task_goal') or '').strip(),
            'task_desc': (item.get('task_desc') or '').strip(),
            'participants': participant_ids,
            'participant_names': [member.name for member in members if member.id in participant_ids],
            'need_meeting': need_meeting,
            'suggested_meeting_date': suggested_meeting_date,
            'milestone_name': milestone.name if milestone else (item.get('milestone') or ''),
            'raw_item': item
        })
    return normalized_items, errors


def create_task_assignment_notification(task):
    if not task.manager:
        return
    content = (
        f'{task.manager.name}您好，您有一个新的待办任务，请查收：\n'
        f'所属项目：{task.project.name if task.project else "-"}\n'
        f'任务名称：{task.name}\n'
        f'开始时间：{task.start_date.strftime("%Y-%m-%d") if task.start_date else "-"}\n'
        f'截止时间：{task.due_date.strftime("%Y-%m-%d") if task.due_date else "-"}\n'
        f'任务目标：{task.objective or "-"}\n'
        f'任务描述：{task.description or "-"}'
    )
    create_notification(
        user_id=task.manager_id,
        title=f'任务分配：{task.name}',
        type='task_assignment',
        project_id=task.project_id,
        related_type='task',
        related_id=task.id,
        content=content
    )


def create_tasks_from_minutes(minutes, normalized_items):
    meeting = minutes.meeting
    created_tasks = []
    ActionItem.query.filter_by(minutes_id=minutes.id).delete(synchronize_session=False)
    for item in normalized_items:
        task = Task(
            project_id=item['project_id'],
            milestone_id=item['milestone_id'],
            name=item['task_name'],
            objective=item['task_goal'],
            description=item['task_desc'],
            manager_id=item['assignee_id'],
            participants=','.join(str(participant_id) for participant_id in item['participants']),
            start_date=item['start_date'],
            due_date=item['end_date'],
            priority=item['priority'],
            status='not_started',
            need_meeting=item['need_meeting'],
            source='minutes',
            meeting_id=meeting.id,
            suggested_meeting_date=item['suggested_meeting_date']
        )
        db.session.add(task)
        db.session.flush()
        action_item = ActionItem(
            minutes_id=minutes.id,
            title=item['task_name'],
            description=item['task_desc'],
            assignee_id=item['assignee_id'],
            participants=','.join(str(participant_id) for participant_id in item['participants']),
            due_date=item['end_date'],
            priority=item['priority'],
            status='confirmed',
            original_text=json.dumps(item['raw_item'], ensure_ascii=False),
            confidence=1.0,
            task_id=task.id
        )
        db.session.add(action_item)
        create_task_assignment_notification(task)
        created_tasks.append({
            'id': task.id,
            'name': task.name,
            'manager_name': item['assignee_name'],
            'milestone_name': item['milestone_name'] or '-',
            'priority': task.priority,
            'start_date': task.start_date.strftime('%Y-%m-%d') if task.start_date else None,
            'due_date': task.due_date.strftime('%Y-%m-%d') if task.due_date else None,
            'need_meeting': task.need_meeting,
            'suggested_meeting_date': task.suggested_meeting_date.strftime('%Y-%m-%d') if task.suggested_meeting_date else None,
            'objective': task.objective,
            'description': task.description
        })
    meeting.tasks_generated = bool(created_tasks)
    minutes.generated_task_count = len(created_tasks)
    return created_tasks


def build_minutes_generated_tasks(minutes_id):
    action_items = ActionItem.query.filter(
        ActionItem.minutes_id == minutes_id,
        ActionItem.task_id.isnot(None)
    ).order_by(ActionItem.id.asc()).all()
    result = []
    for action_item in action_items:
        task = db.session.get(Task, action_item.task_id)
        if not task:
            continue
        result.append({
            'id': task.id,
            'name': task.name,
            'status': task.status,
            'manager_name': task.manager.name if task.manager else '-',
            'milestone_name': task.milestone.name if task.milestone else '-',
            'priority': task.priority,
            'start_date': task.start_date.strftime('%Y-%m-%d') if task.start_date else None,
            'due_date': task.due_date.strftime('%Y-%m-%d') if task.due_date else None,
            'need_meeting': task.need_meeting,
            'suggested_meeting_date': task.suggested_meeting_date.strftime('%Y-%m-%d') if task.suggested_meeting_date else None,
            'objective': task.objective,
            'description': task.description
        })
    return result


def parse_optional_int(value):
    if value in [None, '']:
        return None
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def normalize_user_role(role):
    role_map = {
        'manager': 'project_manager',
        'leader': 'project_owner',
        'owner': 'project_owner',
        'observer': 'member'
    }
    return role_map.get(role, role or 'member')


def normalize_project_role(role):
    role_map = {
        'manager': 'project_manager',
        'leader': 'project_owner',
        'owner': 'project_owner',
        'observer': 'member'
    }
    return role_map.get(role, role or 'member')


def get_request_user():
    data = request.get_json(silent=True) or {}
    user_id = parse_optional_int(request.headers.get('X-User-Id'))
    if user_id is None:
        user_id = parse_optional_int(request.args.get('operator_id'))
    if user_id is None:
        user_id = parse_optional_int(data.get('operator_id'))
    if user_id is None:
        return None
    return User.query.get(user_id)


def require_admin():
    user = get_request_user()
    if not user or normalize_user_role(user.role) != 'admin':
        return None, (jsonify({'success': False, 'message': '仅管理员可执行该操作'}), 403)
    return user, None


def build_project_member_display_role(project, project_member):
    labels = []
    if project.manager_id == project_member.user_id:
        labels.append('项目经理')
    if project.owner_id == project_member.user_id:
        labels.append('项目负责人')
    if not labels:
        role_map = {
            'project_manager': '项目经理',
            'project_owner': '项目负责人',
            'member': '成员'
        }
        labels.append(role_map.get(normalize_project_role(project_member.role), '成员'))
    return '/'.join(labels)


def ensure_project_member(project_id, user_id, role='member'):
    if not project_id or not user_id:
        return None
    normalized_role = normalize_project_role(role)
    pm = ProjectMember.query.filter_by(project_id=project_id, user_id=user_id).first()
    if not pm:
        pm = ProjectMember(project_id=project_id, user_id=user_id, role=normalized_role)
        db.session.add(pm)
    else:
        pm.role = normalized_role
    return pm


def sync_project_core_members(project):
    if project.manager_id:
        ensure_project_member(project.id, project.manager_id, 'project_manager')
    if project.owner_id:
        ensure_project_member(project.id, project.owner_id, 'project_owner')
    for project_member in ProjectMember.query.filter_by(project_id=project.id).all():
        if project_member.user_id == project.manager_id:
            project_member.role = 'project_manager'
        elif project_member.user_id == project.owner_id:
            project_member.role = 'project_owner'
        elif normalize_project_role(project_member.role) in ['project_manager', 'project_owner']:
            project_member.role = 'member'


def enrich_project_item(project, item=None):
    if item is None:
        item = serialize_model(project)
    item['manager_name'] = project.manager.name if project.manager else None
    item['project_manager_name'] = item['manager_name']
    item['owner_name'] = project.owner.name if project.owner else None
    item['project_owner_name'] = item['owner_name']
    return item


def build_recent_meetings(user_id, limit=5):
    if not user_id:
        return []
    organizer_ids = [meeting.id for meeting in Meeting.query.filter_by(organizer_id=user_id).all()]
    attendee_ids = [attendee.meeting_id for attendee in MeetingAttendee.query.filter_by(user_id=user_id).all()]
    related_ids = list(set(organizer_ids + attendee_ids))
    if not related_ids:
        return []
    now = datetime.now()
    meetings = Meeting.query.filter(Meeting.id.in_(related_ids)).all()
    meetings.sort(
        key=lambda meeting: (
            meeting.meeting_date < now if meeting.meeting_date else True,
            meeting.meeting_date if meeting.meeting_date and meeting.meeting_date >= now else datetime.max,
            -(meeting.meeting_date.timestamp()) if meeting.meeting_date and meeting.meeting_date < now else 0
        )
    )
    result = []
    for meeting in meetings[:limit]:
        item = serialize_model(meeting)
        item['organizer_name'] = meeting.organizer.name if meeting.organizer else None
        item['project_name'] = meeting.project.name if meeting.project else None
        item['time_tag'] = '已结束' if get_meeting_display_status(meeting, now) == 'completed' else '未开始'
        result.append(item)
    return result


def get_meeting_display_status(meeting, now=None):
    if now is None:
        now = datetime.now()
    if meeting.status == 'cancelled':
        return 'cancelled'
    if meeting.meeting_date:
        return 'completed' if meeting.meeting_date <= now else 'scheduled'
    return 'completed' if meeting.status == 'completed' else 'scheduled'


def is_admin_user(user):
    return bool(user and normalize_user_role(user.role) == 'admin')


def get_accessible_project_ids(user_id):
    if not user_id:
        return set()
    user = User.query.get(user_id)
    if is_admin_user(user):
        return None
    member_project_ids = [pm.project_id for pm in ProjectMember.query.filter_by(user_id=user_id).all()]
    core_project_ids = [
        p.id for p in Project.query.filter(
            db.or_(Project.manager_id == user_id, Project.owner_id == user_id)
        ).all()
    ]
    return set(member_project_ids + core_project_ids)


def can_user_access_project(user_id, project_id):
    accessible_ids = get_accessible_project_ids(user_id)
    if accessible_ids is None:
        return True
    return project_id in accessible_ids


def can_edit_project(user, project):
    if not user or not project:
        return False
    if is_admin_user(user):
        return True
    return user.id in [project.manager_id, project.owner_id]


def can_edit_task(user, task):
    if not user or not task:
        return False
    if is_admin_user(user):
        return True
    if can_edit_project(user, task.project):
        return True
    return task.manager_id == user.id


@app.before_request
def log_request():
    if request.path.startswith('/api/'):
        print(f"\n=== 请求: {request.method} {request.path} ===")
        if request.is_json:
            print(f"JSON数据: {request.json}")


@app.errorhandler(Exception)
def handle_exception(e):
    import traceback
    print("=== 全局异常 ===")
    print(traceback.format_exc())
    return jsonify({'success': False, 'message': str(e)}), 500


@app.errorhandler(400)
def handle_400(e):
    return jsonify({'success': False, 'message': str(e)}), 400


@app.errorhandler(404)
def handle_404(e):
    return jsonify({'success': False, 'message': '资源不存在'}), 404


@app.route('/api/test', methods=['GET', 'POST'])
def test_api():
    print("=== 测试API被调用 ===")
    return jsonify({'success': True, 'message': 'API正常工作'})


@app.route('/api/login', methods=['POST'])
def login():
    data = request.json
    username = data.get('username')
    password = data.get('password')
    user = User.query.filter_by(username=username, password=password).first()
    if user:
        return jsonify({
            'success': True,
            'data': {
                'id': user.id,
                'username': user.username,
                'name': user.name,
                'role': normalize_user_role(user.role),
                'email': user.email,
                'phone': user.phone
            }
        })
    return jsonify({'success': False, 'message': '用户名或密码错误'})


@app.route('/api/dashboard', methods=['GET'])
def get_dashboard():
    user_id = request.args.get('user_id', type=int)
    current_user = User.query.get(user_id) if user_id else None
    accessible_project_ids = get_accessible_project_ids(user_id)
    total_projects = Project.query.count() if accessible_project_ids is None else len(accessible_project_ids)
    today = datetime.now().date()
    all_user_tasks_query = Task.query
    if accessible_project_ids is not None:
        if not accessible_project_ids:
            all_user_tasks_query = Task.query.filter(Task.id == -1)
        else:
            all_user_tasks_query = all_user_tasks_query.filter(Task.project_id.in_(accessible_project_ids))
    if current_user and normalize_user_role(current_user.role) == 'member':
        all_user_tasks_query = all_user_tasks_query.filter_by(manager_id=user_id)
    all_user_tasks = all_user_tasks_query.all()
    in_progress_count = 0
    overdue_count = 0
    for task in all_user_tasks:
        display_status = get_task_display_status(task, today)
        if display_status == 'overdue':
            overdue_count += 1
        elif display_status == 'in_progress':
            in_progress_count += 1
    pending_meetings = Meeting.query.filter_by(organizer_id=user_id, status='pending').count()
    unread_notifications = Notification.query.filter_by(user_id=user_id, is_read=False).count()
    meeting_reminders = build_meeting_reminders(limit=5, user_id=user_id)
    recent_meetings = build_recent_meetings(user_id, limit=5)
    return jsonify({
        'success': True,
        'data': {
            'my_projects': total_projects,
            'my_tasks': in_progress_count,
            'overdue_tasks': overdue_count,
            'pending_meetings': pending_meetings,
            'unread_notifications': unread_notifications,
            'meeting_reminders': meeting_reminders,
            'recent_meetings': recent_meetings
        }
    })


@app.route('/api/projects', methods=['GET'])
def get_projects():
    page = request.args.get('page', 1, type=int)
    size = request.args.get('size', 10, type=int)
    status = request.args.get('status')
    user_id = request.args.get('user_id', type=int)
    keyword = request.args.get('keyword')
    query = Project.query
    if status:
        query = query.filter_by(status=status)
    if keyword:
        query = query.filter(
            db.or_(
                Project.name.contains(keyword),
                Project.code.contains(keyword)
            )
        )
    if user_id:
        accessible_project_ids = get_accessible_project_ids(user_id)
        if accessible_project_ids is not None:
            if not accessible_project_ids:
                query = query.filter(Project.id == -1)
            else:
                query = query.filter(Project.id.in_(accessible_project_ids))
    total = query.count()
    projects = query.order_by(Project.created_at.desc()).offset((page - 1) * size).limit(size).all()
    result = []
    for p in projects:
        item = enrich_project_item(p)
        item['member_count'] = ProjectMember.query.filter_by(project_id=p.id).count()
        item['task_count'] = Task.query.filter_by(project_id=p.id).count()
        item['completed_task_count'] = Task.query.filter_by(project_id=p.id, status='completed').count()
        item['milestone_count'] = Milestone.query.filter_by(project_id=p.id).count()
        item['completed_milestone_count'] = Milestone.query.filter_by(
            project_id=p.id, status='completed'
        ).count()
        result.append(item)
    return jsonify({
        'success': True,
        'data': {
            'list': result,
            'total': total,
            'page': page,
            'size': size
        }
    })


@app.route('/api/projects', methods=['POST'])
def create_project():
    _, error_response = require_admin()
    if error_response:
        return error_response
    data = request.json
    project = Project(
        name=data.get('name'),
        code=data.get('code'),
        objective=data.get('objective'),
        background=data.get('background'),
        start_date=parse_optional_date(data.get('start_date')),
        end_date=parse_optional_date(data.get('end_date')),
        manager_id=parse_optional_int(data.get('manager_id')),
        owner_id=parse_optional_int(data.get('owner_id')),
        status=data.get('status', 'not_started'),
        description=data.get('description'),
        project_type=data.get('project_type')
    )
    db.session.add(project)
    db.session.commit()
    sync_project_core_members(project)
    if data.get('member_ids'):
        for member_id in data.get('member_ids'):
            ensure_project_member(project.id, parse_optional_int(member_id), 'member')
    db.session.commit()
    return jsonify({'success': True, 'data': {'id': project.id}})


@app.route('/api/projects/<int:project_id>', methods=['GET'])
def get_project(project_id):
    user_id = request.args.get('user_id', type=int)
    project = Project.query.get(project_id)
    if not project:
        return jsonify({'success': False, 'message': '项目不存在'}), 404
    if user_id and not can_user_access_project(user_id, project_id):
        return jsonify({'success': False, 'message': '无权查看该项目'}), 403
    item = enrich_project_item(project)
    item['members'] = []
    for pm in ProjectMember.query.filter_by(project_id=project_id).all():
        member_data = serialize_model(pm.user, exclude=['password'])
        member_data['member_id'] = pm.id
        member_data['role'] = normalize_user_role(pm.user.role)
        member_data['project_role'] = normalize_project_role(pm.role)
        member_data['project_role_display'] = build_project_member_display_role(project, pm)
        member_data['joined_at'] = pm.joined_at.strftime('%Y-%m-%d %H:%M:%S')
        item['members'].append(member_data)
    item['milestones'] = [serialize_model(m) for m in Milestone.query.filter_by(project_id=project_id).all()]
    item['tasks'] = []
    today = datetime.now().date()
    for t in Task.query.filter_by(project_id=project_id).all():
        task_data = serialize_model(t)
        task_data['manager_name'] = t.manager.name if t.manager else None
        task_data['project_manager_id'] = project.manager_id
        task_data['project_owner_id'] = project.owner_id
        task_data['status'] = get_task_display_status(t, today)
        item['tasks'].append(task_data)
    item['meetings'] = []
    for meeting in Meeting.query.filter_by(project_id=project_id).all():
        meeting_data = serialize_model(meeting)
        meeting_data['status'] = get_meeting_display_status(meeting)
        meeting_data['organizer_name'] = meeting.organizer.name if meeting.organizer else None
        item['meetings'].append(meeting_data)
    item['statistics'] = {
        'total_tasks': Task.query.filter_by(project_id=project_id).count(),
        'completed_tasks': Task.query.filter_by(project_id=project_id, status='completed').count(),
        'overdue_tasks': Task.query.filter_by(project_id=project_id).filter(
            Task.status != 'completed', Task.due_date < datetime.now().date()
        ).count(),
        'total_milestones': Milestone.query.filter_by(project_id=project_id).count(),
        'completed_milestones': Milestone.query.filter_by(
            project_id=project_id, status='completed'
        ).count(),
        'total_meetings': Meeting.query.filter_by(project_id=project_id).count(),
        'member_count': ProjectMember.query.filter_by(project_id=project_id).count()
    }
    return jsonify({'success': True, 'data': item})


@app.route('/api/projects/<int:project_id>', methods=['PUT'])
def update_project(project_id):
    operator = get_request_user()
    project = Project.query.get(project_id)
    if not project:
        return jsonify({'success': False, 'message': '项目不存在'}), 404
    if not can_edit_project(operator, project):
        return jsonify({'success': False, 'message': '无权编辑该项目'}), 403
    data = request.json
    for key, value in data.items():
        if hasattr(project, key) and key not in ['id', 'created_at']:
            if key in ['start_date', 'end_date']:
                value = parse_optional_date(value)
            if key in ['manager_id', 'owner_id']:
                value = parse_optional_int(value)
            setattr(project, key, value)
    sync_project_core_members(project)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/projects/<int:project_id>', methods=['DELETE'])
def delete_project(project_id):
    operator = get_request_user()
    project = Project.query.get(project_id)
    if not project:
        return jsonify({'success': False, 'message': '项目不存在'}), 404
    if not can_edit_project(operator, project):
        return jsonify({'success': False, 'message': '无权删除该项目'}), 403
    db.session.delete(project)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/milestones', methods=['GET'])
def get_milestones():
    project_id = request.args.get('project_id', type=int)
    status = request.args.get('status')
    query = Milestone.query
    if project_id:
        query = query.filter_by(project_id=project_id)
    if status:
        query = query.filter_by(status=status)
    milestones = query.order_by(Milestone.planned_date).all()
    result = []
    for m in milestones:
        item = serialize_model(m)
        item['manager_name'] = m.manager.name if m.manager else None
        item['project_name'] = m.project.name if m.project else None
        item['task_count'] = Task.query.filter_by(milestone_id=m.id).count()
        result.append(item)
    return jsonify({'success': True, 'data': result})


@app.route('/api/milestones', methods=['POST'])
def create_milestone():
    data = request.json
    milestone = Milestone(
        project_id=data.get('project_id'),
        name=data.get('name'),
        objective=data.get('objective'),
        planned_date=parse_optional_date(data.get('planned_date')),
        manager_id=parse_optional_int(data.get('manager_id')),
        status=data.get('status', 'not_started'),
        description=data.get('description'),
        need_meeting=data.get('need_meeting', False),
        expected_meeting_date=parse_optional_date(data.get('expected_meeting_date')),
        meeting_participants=data.get('meeting_participants'),
        meeting_objective=data.get('meeting_objective')
    )
    db.session.add(milestone)
    db.session.commit()
    return jsonify({'success': True, 'data': {'id': milestone.id}})


@app.route('/api/milestones/<int:milestone_id>', methods=['GET'])
def get_milestone(milestone_id):
    milestone = Milestone.query.get(milestone_id)
    if not milestone:
        return jsonify({'success': False, 'message': '里程碑不存在'}), 404
    item = serialize_model(milestone)
    item['manager_name'] = milestone.manager.name if milestone.manager else None
    item['project_name'] = milestone.project.name if milestone.project else None
    item['meeting_defaults'] = build_milestone_meeting_defaults(milestone)
    return jsonify({'success': True, 'data': item})


@app.route('/api/milestones/<int:milestone_id>', methods=['PUT'])
def update_milestone(milestone_id):
    milestone = Milestone.query.get(milestone_id)
    if not milestone:
        return jsonify({'success': False, 'message': '里程碑不存在'}), 404
    data = request.json
    for key, value in data.items():
        if hasattr(milestone, key) and key not in ['id', 'created_at']:
            if key in ['planned_date', 'actual_date', 'expected_meeting_date']:
                value = parse_optional_date(value)
            if key in ['project_id', 'manager_id']:
                value = parse_optional_int(value)
            setattr(milestone, key, value)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/milestones/<int:milestone_id>', methods=['DELETE'])
def delete_milestone(milestone_id):
    milestone = Milestone.query.get(milestone_id)
    if not milestone:
        return jsonify({'success': False, 'message': '里程碑不存在'}), 404
    db.session.delete(milestone)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/tasks', methods=['GET'])
def get_tasks():
    page = request.args.get('page', 1, type=int)
    size = request.args.get('size', 10, type=int)
    project_id = request.args.get('project_id', type=int)
    manager_id = request.args.get('manager_id', type=int)
    user_id = request.args.get('user_id', type=int)
    status = request.args.get('status')
    need_meeting = request.args.get('need_meeting', type=str)
    query = Task.query
    accessible_project_ids = None
    if user_id:
        accessible_project_ids = get_accessible_project_ids(user_id)
        if accessible_project_ids is not None:
            if not accessible_project_ids:
                return jsonify({
                    'success': True,
                    'data': {'list': [], 'total': 0, 'page': page, 'size': size}
                })
            query = query.filter(Task.project_id.in_(accessible_project_ids))
    if project_id:
        if user_id and not can_user_access_project(user_id, project_id):
            return jsonify({
                'success': False,
                'message': '无权查看该项目任务'
            }), 403
        query = query.filter_by(project_id=project_id)
    if manager_id:
        query = query.filter_by(manager_id=manager_id)
    if status == 'completed':
        query = query.filter_by(status='completed')
    elif status:
        query = query.filter(Task.status != 'completed')
    if need_meeting == 'true':
        query = query.filter_by(need_meeting=True, meeting_invitation_created=False)
    all_tasks = query.order_by(Task.created_at.desc()).all()
    today = datetime.now().date()
    result = []
    for t in all_tasks:
        item = serialize_model(t)
        item['manager_name'] = t.manager.name if t.manager else None
        item['project_name'] = t.project.name if t.project else None
        item['milestone_name'] = t.milestone.name if t.milestone else None
        item['project_manager_id'] = t.project.manager_id if t.project else None
        item['project_owner_id'] = t.project.owner_id if t.project else None
        item['status'] = get_task_display_status(t, today)
        if status and item['status'] != status:
            continue
        result.append(item)
    total = len(result)
    start = (page - 1) * size
    end = start + size
    paged_result = result[start:end]
    return jsonify({
        'success': True,
        'data': {
            'list': paged_result,
            'total': total,
            'page': page,
            'size': size
        }
    })


@app.route('/api/tasks', methods=['POST'])
def create_task():
    operator = get_request_user()
    data = request.json
    project_id = parse_optional_int(data.get('project_id'))
    if not operator:
        return jsonify({'success': False, 'message': '请先登录后再创建任务'}), 403
    if not project_id or not can_user_access_project(operator.id, project_id):
        return jsonify({'success': False, 'message': '无权在该项目下创建任务'}), 403
    task = Task(
        project_id=project_id,
        milestone_id=data.get('milestone_id'),
        name=data.get('name'),
        objective=data.get('objective'),
        description=data.get('description'),
        manager_id=data.get('manager_id'),
        participants=data.get('participants'),
        start_date=parse_optional_date(data.get('start_date')),
        due_date=parse_optional_date(data.get('due_date')),
        priority=data.get('priority', 'medium'),
        status=data.get('status', 'not_started'),
        need_meeting=data.get('need_meeting', False),
        meeting_objective=None,
        suggested_participants=None,
        source=data.get('source', 'manual'),
        suggested_meeting_date=parse_optional_date(data.get('suggested_meeting_date')),
        suggested_meeting_topic=None
    )
    db.session.add(task)
    db.session.commit()
    return jsonify({'success': True, 'data': {'id': task.id}})


@app.route('/api/tasks/<int:task_id>', methods=['GET'])
def get_task(task_id):
    user_id = request.args.get('user_id', type=int)
    task = Task.query.get(task_id)
    if not task:
        return jsonify({'success': False, 'message': '任务不存在'}), 404
    if user_id and not can_user_access_project(user_id, task.project_id):
        return jsonify({'success': False, 'message': '无权查看该任务'}), 403
    item = serialize_model(task)
    item['manager_name'] = task.manager.name if task.manager else None
    item['project_name'] = task.project.name if task.project else None
    item['milestone_name'] = task.milestone.name if task.milestone else None
    item['project_manager_id'] = task.project.manager_id if task.project else None
    item['project_owner_id'] = task.project.owner_id if task.project else None
    item['meeting_defaults'] = build_task_meeting_defaults(task)
    if task.source_meeting:
        item['source_meeting_title'] = task.source_meeting.title
    return jsonify({'success': True, 'data': item})


@app.route('/api/tasks/<int:task_id>', methods=['PUT'])
def update_task(task_id):
    operator = get_request_user()
    task = Task.query.get(task_id)
    if not task:
        return jsonify({'success': False, 'message': '任务不存在'}), 404
    if not can_edit_task(operator, task):
        return jsonify({'success': False, 'message': '无权编辑该任务'}), 403
    data = request.json
    for key, value in data.items():
        if hasattr(task, key) and key not in ['id', 'created_at']:
            if key in ['start_date', 'due_date', 'suggested_meeting_date']:
                value = parse_optional_date(value)
            setattr(task, key, value)
    task.meeting_objective = None
    task.suggested_participants = None
    task.suggested_meeting_topic = None
    if not task.need_meeting:
        task.suggested_meeting_date = None
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/tasks/<int:task_id>/complete', methods=['POST'])
def complete_task(task_id):
    operator = get_request_user()
    task = Task.query.get(task_id)
    if not task:
        return jsonify({'success': False, 'message': '任务不存在'}), 404
    if not can_edit_task(operator, task):
        return jsonify({'success': False, 'message': '无权操作该任务'}), 403
    task.status = 'completed'
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/tasks/<int:task_id>', methods=['DELETE'])
def delete_task(task_id):
    operator = get_request_user()
    task = Task.query.get(task_id)
    if not task:
        return jsonify({'success': False, 'message': '任务不存在'}), 404
    if not can_edit_task(operator, task):
        return jsonify({'success': False, 'message': '无权删除该任务'}), 403
    db.session.delete(task)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/meetings', methods=['GET'])
def get_meetings():
    page = request.args.get('page', 1, type=int)
    size = request.args.get('size', 10, type=int)
    project_id = request.args.get('project_id', type=int)
    status = request.args.get('status')
    query = Meeting.query
    if project_id:
        query = query.filter_by(project_id=project_id)
    meetings = query.order_by(Meeting.meeting_date.desc()).all()
    now = datetime.now()
    result = []
    for m in meetings:
        display_status = get_meeting_display_status(m, now)
        if status and display_status != status:
            continue
        item = serialize_model(m)
        item['status'] = display_status
        item['organizer_name'] = m.organizer.name if m.organizer else None
        item['project_name'] = m.project.name if m.project else None
        item['attendee_count'] = MeetingAttendee.query.filter_by(meeting_id=m.id).count()
        result.append(item)
    total = len(result)
    start = (page - 1) * size
    end = start + size
    paged_result = result[start:end]
    return jsonify({
        'success': True,
        'data': {
            'list': paged_result,
            'total': total,
            'page': page,
            'size': size
        }
    })


@app.route('/api/meetings/reminders', methods=['GET'])
def get_meeting_reminders():
    user_id = request.args.get('user_id', type=int)
    return jsonify({'success': True, 'data': build_meeting_reminders(user_id=user_id)})


@app.route('/api/meetings', methods=['POST'])
def create_meeting():
    data = request.json
    print("=== 创建会议请求数据 ===")
    print(data)
    try:
        if not data.get('project_id'):
            return jsonify({'success': False, 'message': '请选择项目'}), 400
        if not data.get('title'):
            return jsonify({'success': False, 'message': '请填写会议主题'}), 400
        if not data.get('meeting_date'):
            return jsonify({'success': False, 'message': '请选择会议时间'}), 400
        date_str = data.get('meeting_date')
        print(f"原始日期字符串: {date_str}")
        meeting_date = parse_meeting_datetime(date_str)
        print(f"解析后的会议日期: {meeting_date}")
        meeting = Meeting(
            project_id=data.get('project_id'),
            title=data.get('title'),
            objective=data.get('objective'),
            meeting_date=meeting_date,
            location=data.get('location'),
            meeting_type=data.get('meeting_type'),
            organizer_id=data.get('organizer_id'),
            description=data.get('description'),
            source=data.get('source', 'manual'),
            source_task_id=data.get('source_task_id'),
            source_milestone_id=data.get('source_milestone_id'),
            status=data.get('status', 'scheduled')
        )
        db.session.add(meeting)
        db.session.commit()
        print(f"会议创建成功, ID: {meeting.id}")
        attendee_ids = parse_id_list(data.get('attendee_ids'))
        if attendee_ids:
            for attendee_id in attendee_ids:
                ma = MeetingAttendee(meeting_id=meeting.id, user_id=attendee_id)
                db.session.add(ma)
            create_meeting_invitation_notifications(meeting, attendee_ids)
            db.session.commit()
        if data.get('source_task_id'):
            task = Task.query.get(data.get('source_task_id'))
            if task:
                task.meeting_invitation_created = True
                task.meeting_id = meeting.id
                db.session.commit()
        if data.get('source_milestone_id'):
            pass
        return jsonify({'success': True, 'data': {'id': meeting.id}})
    except Exception as e:
        import traceback
        print("=== 创建会议出错 ===")
        print(traceback.format_exc())
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/api/meetings/<int:meeting_id>', methods=['GET'])
def get_meeting(meeting_id):
    meeting = Meeting.query.get(meeting_id)
    if not meeting:
        return jsonify({'success': False, 'message': '会议不存在'}), 404
    item = serialize_model(meeting)
    item['status'] = get_meeting_display_status(meeting)
    item['organizer_name'] = meeting.organizer.name if meeting.organizer else None
    item['project_name'] = meeting.project.name if meeting.project else None
    item['attendees'] = []
    for ma in MeetingAttendee.query.filter_by(meeting_id=meeting_id).all():
        attendee_data = serialize_model(ma.user, exclude=['password'])
        attendee_data['is_required'] = ma.is_required
        item['attendees'].append(attendee_data)
    minutes = MeetingMinutes.query.filter_by(meeting_id=meeting_id).order_by(MeetingMinutes.created_at.desc()).first()
    if minutes:
        item['minutes'] = serialize_model(minutes)
        item['minutes']['action_items'] = []
        for ai in ActionItem.query.filter_by(minutes_id=minutes.id).all():
            action_data = serialize_model(ai)
            action_data['assignee_name'] = ai.assignee.name if ai.assignee else None
            item['minutes']['action_items'].append(action_data)
        item['generated_tasks'] = build_minutes_generated_tasks(minutes.id)
    else:
        item['minutes'] = None
        item['generated_tasks'] = []
    return jsonify({'success': True, 'data': item})


@app.route('/api/meetings/<int:meeting_id>', methods=['PUT'])
def update_meeting(meeting_id):
    meeting = Meeting.query.get(meeting_id)
    if not meeting:
        return jsonify({'success': False, 'message': '会议不存在'}), 404
    data = request.json
    for key, value in data.items():
        if hasattr(meeting, key) and key not in ['id', 'created_at']:
            if key == 'meeting_date' and value:
                value = parse_meeting_datetime(value)
            setattr(meeting, key, value)
    if data.get('attendee_ids'):
        MeetingAttendee.query.filter_by(meeting_id=meeting_id).delete()
        for attendee_id in data.get('attendee_ids'):
            ma = MeetingAttendee(meeting_id=meeting.id, user_id=attendee_id)
            db.session.add(ma)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/meetings/<int:meeting_id>', methods=['DELETE'])
def delete_meeting(meeting_id):
    meeting = Meeting.query.get(meeting_id)
    if not meeting:
        return jsonify({'success': False, 'message': '会议不存在'}), 404
    MeetingAttendee.query.filter_by(meeting_id=meeting_id).delete()
    minutes_list = MeetingMinutes.query.filter_by(meeting_id=meeting_id).all()
    for minutes in minutes_list:
        ActionItem.query.filter_by(minutes_id=minutes.id).delete()
    MeetingMinutes.query.filter_by(meeting_id=meeting_id).delete()
    db.session.delete(meeting)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/minutes/upload', methods=['POST'])
def upload_minutes():
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': '没有上传文件'}), 400
    file = request.files['file']
    meeting_id = request.form.get('meeting_id', type=int)
    if not meeting_id:
        return jsonify({'success': False, 'message': '缺少会议ID'}), 400
    meeting = Meeting.query.get(meeting_id)
    if not meeting:
        return jsonify({'success': False, 'message': '会议不存在'}), 404
    if get_meeting_display_status(meeting) != 'completed':
        return jsonify({'success': False, 'message': '会议未结束，暂不能上传会议纪要'}), 400
    filename = secure_filename(file.filename)
    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
    filename = f'{timestamp}_{filename}'
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)
    manual_content = request.form.get('manual_content', '')
    try:
        content = clean_text_content(extract_minutes_text(file_path, file.filename))
    except Exception as exc:
        if not manual_content.strip():
            return jsonify({'success': False, 'message': f'纪要文件读取失败：{exc}'}), 400
        content = ''
    if manual_content.strip():
        content = clean_text_content('\n\n'.join(filter(None, [content, manual_content])))
    minutes = MeetingMinutes(
        meeting_id=meeting_id,
        file_path=file_path,
        file_name=file.filename,
        content=content,
        parse_status='uploaded',
        parse_error=None,
        generated_task_count=0
    )
    db.session.add(minutes)
    meeting.minutes_uploaded = True
    db.session.commit()
    return jsonify({'success': True, 'data': {'id': minutes.id, 'parse_status': minutes.parse_status}})


@app.route('/api/minutes/<int:minutes_id>/parse', methods=['POST'])
def parse_minutes(minutes_id):
    minutes = MeetingMinutes.query.get(minutes_id)
    if not minutes:
        return jsonify({'success': False, 'message': '会议纪要不存在'}), 404
    meeting = minutes.meeting
    existing_generated_tasks = build_minutes_generated_tasks(minutes.id)
    if minutes.parse_status == 'parsed':
        return jsonify({
            'success': True,
            'data': {
                'minutes_id': minutes.id,
                'parse_status': minutes.parse_status or 'parsed',
                'task_count': len(existing_generated_tasks),
                'tasks': existing_generated_tasks,
                'errors': minutes.parse_error.splitlines() if minutes.parse_error else [],
                'message': '该上传的会议纪要已解析过，已返回当前文件生成的任务结果'
            }
        })
    content = clean_text_content(minutes.content or '')
    if not content:
        return jsonify({'success': False, 'message': '会议纪要内容为空，无法解析'}), 400
    minutes.parse_status = 'parsing'
    minutes.parse_error = None
    db.session.commit()
    try:
        prompt_text, context = build_minutes_task_prompt(meeting, content)
        llm_response = call_minutes_llm(prompt_text)
        llm_items = extract_json_array(llm_response)
        meeting_date = meeting.meeting_date.date() if meeting.meeting_date else datetime.now().date()
        normalized_items, errors = normalize_minutes_task_items(llm_items, context, meeting_date)
        created_tasks = create_tasks_from_minutes(minutes, normalized_items)
        minutes.parsed_at = datetime.now()
        minutes.parse_status = 'parsed'
        minutes.parse_error = None if not errors else '\n'.join(errors)
        minutes.generated_task_count = len(created_tasks)
        db.session.commit()
        return jsonify({
            'success': True,
            'data': {
                'minutes_id': minutes.id,
                'parse_status': minutes.parse_status,
                'task_count': len(created_tasks),
                'tasks': created_tasks,
                'errors': errors
            }
        })
    except Exception as exc:
        db.session.rollback()
        minutes = MeetingMinutes.query.get(minutes_id)
        if minutes:
            minutes.parse_status = 'failed'
            minutes.parse_error = str(exc)
            db.session.commit()
        return jsonify({'success': False, 'message': f'纪要解析失败：{exc}'}), 500


def parse_action_items(content, meeting_id):
    action_items = []
    lines = content.split('\n')
    patterns = [
        r'(\d+[\.\、])\s*(.+)',
        r'[-•]\s*(.+)',
        r'【(.+)】',
        r'待办[：:]\s*(.+)',
        r'行动项[：:]\s*(.+)',
        r'责任人[：:]\s*(.+)',
    ]
    current_item = None
    for line in lines:
        line = line.strip()
        if not line:
            continue
        for pattern in patterns:
            match = re.search(pattern, line)
            if match:
                if current_item:
                    action_items.append(current_item)
                title = match.group(1) if match.lastindex else line
                current_item = {
                    'title': title.strip() if isinstance(title, str) else line,
                    'original_text': line,
                    'confidence': 0.7
                }
                assignee_match = re.search(r'[负责|责任人|执行人][：:]\s*(\S+)', line)
                if assignee_match:
                    assignee_name = assignee_match.group(1)
                    user = User.query.filter(User.name.contains(assignee_name)).first()
                    if user:
                        current_item['assignee_id'] = user.id
                date_match = re.search(r'(\d{4}[-/年]\d{1,2}[-/月]\d{1,2}[日]?)', line)
                if date_match:
                    date_str = date_match.group(1)
                    date_str = date_str.replace('年', '-').replace('月', '-').replace('日', '').replace('/', '-')
                    try:
                        current_item['due_date'] = datetime.strptime(date_str, '%Y-%m-%d').date()
                    except:
                        pass
                break
        else:
            if current_item:
                current_item['description'] = current_item.get('description', '') + line + '\n'
    if current_item:
        action_items.append(current_item)
    return action_items


@app.route('/api/action-items/<int:action_id>/confirm', methods=['POST'])
def confirm_action_item(action_id):
    action_item = ActionItem.query.get(action_id)
    if not action_item:
        return jsonify({'success': False, 'message': '行动项不存在'}), 404
    data = request.json
    task = Task(
        project_id=data.get('project_id') or action_item.minutes.meeting.project_id,
        name=action_item.title,
        description=action_item.description or action_item.original_text,
        manager_id=data.get('assignee_id') or action_item.assignee_id,
        due_date=action_item.due_date,
        priority=action_item.priority,
        status='not_started',
        source='minutes',
        meeting_id=action_item.minutes.meeting_id
    )
    db.session.add(task)
    action_item.task_id = task.id
    action_item.status = 'confirmed'
    db.session.commit()
    return jsonify({'success': True, 'data': {'task_id': task.id}})


@app.route('/api/notifications', methods=['GET'])
def get_notifications():
    user_id = request.args.get('user_id', type=int)
    page = request.args.get('page', 1, type=int)
    size = request.args.get('size', 20, type=int)
    type_filter = request.args.get('type')
    is_read = request.args.get('is_read')
    query = Notification.query
    if user_id:
        query = query.filter_by(user_id=user_id)
    if type_filter:
        query = query.filter_by(type=type_filter)
    if is_read is not None:
        query = query.filter_by(is_read=is_read == 'true')
    total = query.count()
    notifications = query.order_by(Notification.created_at.desc()).offset(
        (page - 1) * size
    ).limit(size).all()
    result = []
    for n in notifications:
        item = serialize_model(n)
        item['project_name'] = n.project.name if n.project else None
        result.append(item)
    return jsonify({
        'success': True,
        'data': {
            'list': result,
            'total': total,
            'page': page,
            'size': size
        }
    })


@app.route('/api/notifications/<int:notification_id>/read', methods=['POST'])
def mark_notification_read(notification_id):
    notification = Notification.query.get(notification_id)
    if not notification:
        return jsonify({'success': False, 'message': '通知不存在'}), 404
    notification.is_read = True
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/notifications/read-all', methods=['POST'])
def mark_all_notifications_read():
    user_id = request.json.get('user_id')
    if user_id:
        Notification.query.filter_by(user_id=user_id, is_read=False).update({'is_read': True})
        db.session.commit()
    return jsonify({'success': True})


@app.route('/api/users', methods=['GET'])
def get_users():
    users = User.query.all()
    result = []
    for u in users:
        item = serialize_model(u, exclude=['password'])
        item['role'] = normalize_user_role(u.role)
        result.append(item)
    return jsonify({'success': True, 'data': result})


@app.route('/api/users', methods=['POST'])
def create_user():
    _, error_response = require_admin()
    if error_response:
        return error_response
    data = request.json
    user = User(
        username=data.get('username'),
        password=data.get('password'),
        name=data.get('name'),
        email=data.get('email'),
        phone=data.get('phone'),
        role=normalize_user_role(data.get('role', 'member')),
        can_attend_meeting=data.get('can_attend_meeting', True),
        can_assign_task=data.get('can_assign_task', True)
    )
    db.session.add(user)
    db.session.commit()
    return jsonify({'success': True, 'data': {'id': user.id}})


@app.route('/api/users/<int:user_id>', methods=['PUT'])
def update_user(user_id):
    _, error_response = require_admin()
    if error_response:
        return error_response
    user = User.query.get(user_id)
    if not user:
        return jsonify({'success': False, 'message': '用户不存在'}), 404
    data = request.json
    for key, value in data.items():
        if hasattr(user, key) and key not in ['id', 'created_at']:
            if key == 'role':
                value = normalize_user_role(value)
            setattr(user, key, value)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/users/<int:user_id>', methods=['DELETE'])
def delete_user(user_id):
    operator, error_response = require_admin()
    if error_response:
        return error_response
    user = User.query.get(user_id)
    if not user:
        return jsonify({'success': False, 'message': '用户不存在'}), 404
    if operator.id == user_id:
        return jsonify({'success': False, 'message': '不能删除当前登录的管理员账号'}), 400
    has_related_data = any([
        Project.query.filter(db.or_(Project.manager_id == user_id, Project.owner_id == user_id)).first(),
        ProjectMember.query.filter_by(user_id=user_id).first(),
        Milestone.query.filter_by(manager_id=user_id).first(),
        Task.query.filter_by(manager_id=user_id).first(),
        Meeting.query.filter_by(organizer_id=user_id).first(),
        MeetingAttendee.query.filter_by(user_id=user_id).first(),
        ActionItem.query.filter_by(assignee_id=user_id).first(),
        Notification.query.filter_by(user_id=user_id).first()
    ])
    if has_related_data:
        return jsonify({'success': False, 'message': '该账号存在关联业务数据，无法删除'}), 400
    db.session.delete(user)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/project-members', methods=['GET'])
def get_project_members():
    project_id = request.args.get('project_id', type=int)
    if not project_id:
        return jsonify({'success': False, 'message': '缺少项目ID'}), 400
    project = Project.query.get(project_id)
    if not project:
        return jsonify({'success': False, 'message': '项目不存在'}), 404
    members = ProjectMember.query.filter_by(project_id=project_id).all()
    result = []
    for pm in members:
        item = serialize_model(pm.user, exclude=['password'])
        item['role'] = normalize_user_role(pm.user.role)
        item['project_role'] = normalize_project_role(pm.role)
        item['project_role_display'] = build_project_member_display_role(project, pm)
        item['joined_at'] = pm.joined_at.strftime('%Y-%m-%d %H:%M:%S')
        result.append(item)
    return jsonify({'success': True, 'data': result})


@app.route('/api/project-members', methods=['POST'])
def add_project_member():
    data = request.json
    project_id = parse_optional_int(data.get('project_id'))
    user_id = parse_optional_int(data.get('user_id'))
    role = normalize_project_role(data.get('role', 'member'))
    project = Project.query.get(project_id)
    if not project:
        return jsonify({'success': False, 'message': '项目不存在'}), 404
    if not user_id:
        return jsonify({'success': False, 'message': '缺少成员ID'}), 400
    pm = ensure_project_member(project_id, user_id, role)
    if role == 'project_manager':
        project.manager_id = user_id
    elif role == 'project_owner':
        project.owner_id = user_id
    db.session.commit()
    return jsonify({'success': True, 'data': {'member_id': pm.id}})


@app.route('/api/project-members/<int:member_id>', methods=['DELETE'])
def remove_project_member(member_id):
    pm = ProjectMember.query.get(member_id)
    if not pm:
        return jsonify({'success': False, 'message': '成员关系不存在'}), 404
    project = Project.query.get(pm.project_id)
    if project:
        if project.manager_id == pm.user_id:
            project.manager_id = None
        if project.owner_id == pm.user_id:
            project.owner_id = None
    db.session.delete(pm)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/rules', methods=['GET'])
def get_rules():
    rules = RuleConfig.query.filter_by(is_active=True).all()
    result = [serialize_model(r) for r in rules]
    return jsonify({'success': True, 'data': result})


@app.route('/api/rules/<int:rule_id>', methods=['PUT'])
def update_rule(rule_id):
    rule = RuleConfig.query.get(rule_id)
    if not rule:
        return jsonify({'success': False, 'message': '规则不存在'}), 404
    data = request.json
    for key, value in data.items():
        if hasattr(rule, key) and key not in ['id', 'created_at']:
            setattr(rule, key, value)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/generate-reminders', methods=['POST'])
def generate_reminders():
    tasks = Task.query.filter_by(need_meeting=True, meeting_reminder_created=False).all()
    for task in tasks:
        project = Project.query.get(task.project_id)
        if project and project.manager_id:
            meeting_defaults = build_task_meeting_defaults(task)
            create_notification(
                user_id=project.manager_id,
                title=f'待创建会议提醒：{task.name}',
                type='meeting_reminder',
                project_id=task.project_id,
                related_type='task',
                related_id=task.id,
                content=f'任务"{task.name}"标记需要召开会议，但尚未创建会议邀请。默认会议主题：{meeting_defaults["title"]}，默认时间：{meeting_defaults["meeting_date_display"]}'
            )
            task.meeting_reminder_created = True
    overdue_tasks = Task.query.filter(
        Task.status.in_(['not_started', 'in_progress']),
        Task.due_date < datetime.now().date()
    ).all()
    for task in overdue_tasks:
        project = Project.query.get(task.project_id)
        if project and project.manager_id:
            existing = Notification.query.filter_by(
                user_id=project.manager_id,
                type='task_overdue',
                related_type='task',
                related_id=task.id,
                is_read=False
            ).first()
            if not existing:
                create_notification(
                    user_id=project.manager_id,
                    title=f'任务逾期提醒：{task.name}',
                    type='task_overdue',
                    project_id=task.project_id,
                    related_type='task',
                    related_id=task.id,
                    content=f'任务"{task.name}"已逾期，截止日期：{task.due_date}'
                )
    db.session.commit()
    return jsonify({'success': True, 'message': '提醒生成完成'})


@app.route('/<path:path>')
def serve_static(path):
    return send_from_directory('.', path)


@app.route('/')
def index():
    return send_from_directory('.', 'login.html')


def migrate_legacy_schema():
    with db.engine.begin() as conn:
        project_columns = {row[1] for row in conn.exec_driver_sql("PRAGMA table_info(project)").fetchall()}
        if 'owner_id' not in project_columns:
            conn.exec_driver_sql("ALTER TABLE project ADD COLUMN owner_id INTEGER")
        conn.exec_driver_sql("UPDATE project SET owner_id = manager_id WHERE owner_id IS NULL")
        minutes_columns = {row[1] for row in conn.exec_driver_sql("PRAGMA table_info(meeting_minutes)").fetchall()}
        if 'parse_status' not in minutes_columns:
            conn.exec_driver_sql("ALTER TABLE meeting_minutes ADD COLUMN parse_status VARCHAR(20)")
            conn.exec_driver_sql("UPDATE meeting_minutes SET parse_status = 'uploaded' WHERE parse_status IS NULL")
        if 'parse_error' not in minutes_columns:
            conn.exec_driver_sql("ALTER TABLE meeting_minutes ADD COLUMN parse_error TEXT")
        if 'generated_task_count' not in minutes_columns:
            conn.exec_driver_sql("ALTER TABLE meeting_minutes ADD COLUMN generated_task_count INTEGER DEFAULT 0")
        conn.exec_driver_sql("UPDATE meeting_minutes SET parse_status = 'parsed' WHERE parsed_at IS NOT NULL AND (parse_status IS NULL OR parse_status = 'uploaded')")
        conn.exec_driver_sql("UPDATE user SET role = 'admin' WHERE username = 'admin' AND role = 'manager'")
        conn.exec_driver_sql("UPDATE user SET role = 'project_manager' WHERE role = 'manager'")
        conn.exec_driver_sql("UPDATE user SET role = 'project_owner' WHERE role IN ('leader', 'owner')")
        conn.exec_driver_sql("UPDATE user SET role = 'member' WHERE role = 'observer'")
        conn.exec_driver_sql("UPDATE project_member SET role = 'project_manager' WHERE role = 'manager'")
        conn.exec_driver_sql("UPDATE project_member SET role = 'project_owner' WHERE role IN ('leader', 'owner')")
        conn.exec_driver_sql("UPDATE project_member SET role = 'member' WHERE role = 'observer'")
        conn.exec_driver_sql("UPDATE meeting SET status = 'scheduled' WHERE status = 'pending'")
        existing_rule_keys = {row[0] for row in conn.exec_driver_sql("SELECT rule_key FROM rule_config").fetchall()}
        default_rules = [
            ('enable_email_notification', '启用邮件通知', 'false', '是否在站内提醒之外同步发送邮件通知'),
            ('enable_feishu_notification', '启用飞书机器人通知', 'false', '是否在站内提醒之外同步发送飞书机器人群消息'),
            ('feishu_webhook_url', '飞书 Webhook 地址', '', '飞书机器人 Webhook 地址'),
            ('smtp_server', 'SMTP 服务器', '', '邮件发送的 SMTP 服务器地址'),
            ('smtp_port', 'SMTP 端口', '994', '邮件发送的 SMTP 端口'),
            ('smtp_sender', 'SMTP 发件邮箱', '', '邮件发送使用的发件邮箱'),
            ('smtp_password', 'SMTP 授权码', '', '邮件发送使用的 SMTP 授权码'),
            ('minutes_llm_chat_url', '纪要解析模型地址', 'http://192.168.200.54:11434/api/chat', '会议纪要解析调用的大模型接口地址'),
            ('minutes_llm_model', '纪要解析模型', 'glm-5:cloud', '会议纪要解析使用的大模型名称'),
            ('minutes_llm_temperature', '纪要解析温度', '0.2', '会议纪要解析的大模型温度参数'),
            ('minutes_llm_timeout_seconds', '纪要解析超时秒数', '180', '会议纪要解析调用大模型的超时时间')
        ]
        for rule_key, rule_name, rule_value, description in default_rules:
            if rule_key not in existing_rule_keys:
                conn.exec_driver_sql(
                    "INSERT INTO rule_config (rule_key, rule_name, rule_value, description, is_active, created_at, updated_at) VALUES (?, ?, ?, ?, 1, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)",
                    (rule_key, rule_name, rule_value, description)
                )


def backfill_project_core_members():
    for project in Project.query.all():
        sync_project_core_members(project)
    db.session.commit()


def init_default_data():
    if User.query.count() == 0:
        admin = User(
            username='admin',
            password='admin123',
            name='管理员',
            email='admin@example.com',
            role='admin'
        )
        db.session.add(admin)
        user1 = User(
            username='zhangsan',
            password='123456',
            name='张三',
            email='zhangsan@example.com',
            role='member'
        )
        db.session.add(user1)
        user2 = User(
            username='lisi',
            password='123456',
            name='李四',
            email='lisi@example.com',
            role='member'
        )
        db.session.add(user2)
        db.session.commit()
    if RuleConfig.query.count() == 0:
        rules = [
            RuleConfig(rule_key='milestone_reminder_days', rule_name='里程碑提前提醒天数', rule_value='7', description='里程碑提前多少天触发会议提醒'),
            RuleConfig(rule_key='meeting_minutes_reminder_hours', rule_name='会议纪要提醒小时数', rule_value='24', description='会议结束后多少小时提醒上传纪要'),
            RuleConfig(rule_key='auto_generate_task', rule_name='纪要解析后自动生成任务', rule_value='true', description='纪要解析后是否自动生成任务，否则需人工确认'),
            RuleConfig(rule_key='enable_email_notification', rule_name='启用邮件通知', rule_value='false', description='是否在站内提醒之外同步发送邮件通知'),
            RuleConfig(rule_key='enable_feishu_notification', rule_name='启用飞书机器人通知', rule_value='false', description='是否在站内提醒之外同步发送飞书机器人群消息'),
            RuleConfig(rule_key='feishu_webhook_url', rule_name='飞书 Webhook 地址', rule_value='', description='飞书机器人 Webhook 地址'),
            RuleConfig(rule_key='smtp_server', rule_name='SMTP 服务器', rule_value='', description='邮件发送的 SMTP 服务器地址'),
            RuleConfig(rule_key='smtp_port', rule_name='SMTP 端口', rule_value='994', description='邮件发送的 SMTP 端口'),
            RuleConfig(rule_key='smtp_sender', rule_name='SMTP 发件邮箱', rule_value='', description='邮件发送使用的发件邮箱'),
            RuleConfig(rule_key='smtp_password', rule_name='SMTP 授权码', rule_value='', description='邮件发送使用的 SMTP 授权码'),
            RuleConfig(rule_key='minutes_llm_chat_url', rule_name='纪要解析模型地址', rule_value='http://192.168.200.54:11434/api/chat', description='会议纪要解析调用的大模型接口地址'),
            RuleConfig(rule_key='minutes_llm_model', rule_name='纪要解析模型', rule_value='glm-5:cloud', description='会议纪要解析使用的大模型名称'),
            RuleConfig(rule_key='minutes_llm_temperature', rule_name='纪要解析温度', rule_value='0.2', description='会议纪要解析的大模型温度参数'),
            RuleConfig(rule_key='minutes_llm_timeout_seconds', rule_name='纪要解析超时秒数', rule_value='180', description='会议纪要解析调用大模型的超时时间'),
        ]
        for rule in rules:
            db.session.add(rule)
        db.session.commit()


with app.app_context():
    db.create_all()
    migrate_legacy_schema()
    init_default_data()
    backfill_project_core_members()


if __name__ == '__main__':
    app.run(debug=True, host='127.0.0.1', port=5000, use_reloader=False)
